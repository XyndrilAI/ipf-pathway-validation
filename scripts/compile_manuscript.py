#!/usr/bin/env python3
"""Compile manuscript sections into single Word document for BMC Bioinformatics."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def read_md(filepath: Path) -> str:
    return filepath.read_text(encoding="utf-8")


def add_title_page(doc: Document) -> None:
    """Title, author, affiliation, contact."""
    title = doc.add_heading(
        "Pathway-level IPF signatures demonstrate robust reproducibility "
        "across independent bulk cohorts under strict leakage control",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for text, size in [
        ("Tony Keltakangas", 12),
        ("Software Developer / Independent Researcher", 10),
        ("Fusion Dev Group, Finland", 10),
        ("Correspondence: fusion@xyndril.dev", 10),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)

    doc.add_page_break()


def add_abstract(doc: Document, paper_dir: Path) -> None:
    """Parse structured abstract with bold section labels."""
    doc.add_heading("Abstract", level=1)

    raw = read_md(paper_dir / "abstract.md")
    # Remove the '# Abstract' header line
    raw = re.sub(r"^#\s+Abstract\s*\n", "", raw, count=1)
    # Remove the Keywords line — we add it separately
    keywords_match = re.search(r"\*\*Keywords:\*\*\s*(.+)", raw)
    keywords_text = keywords_match.group(1).strip() if keywords_match else ""
    raw = re.sub(r"\*\*Keywords:\*\*.*", "", raw)

    # Split by structured labels: **Background:** etc.
    sections = re.split(r"\*\*(\w+):\*\*", raw.strip())
    # sections[0] is before first label (empty), then alternating label/content
    i = 1
    while i < len(sections) - 1:
        label = sections[i].strip()
        content = sections[i + 1].strip()
        # Collapse newlines within a section into single spaces
        content = re.sub(r"\s*\n\s*", " ", content)
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(content)
        i += 2

    # Keywords
    if keywords_text:
        p = doc.add_paragraph()
        run_kw = p.add_run("Keywords: ")
        run_kw.bold = True
        p.add_run(keywords_text)

    doc.add_page_break()


def add_markdown_section(doc: Document, title: str, filepath: Path, *, page_break: bool = True) -> None:
    """Add a full markdown section, handling ## sub-headings and **bold**."""
    doc.add_heading(title, level=1)

    raw = read_md(filepath)
    # Remove the top-level '# Title' line
    raw = re.sub(r"^#\s+[^\n]+\n", "", raw, count=1)

    lines = raw.split("\n")
    current_para_lines: list[str] = []

    def flush_para() -> None:
        text = " ".join(current_para_lines).strip()
        current_para_lines.clear()
        if not text:
            return
        # Handle inline bold: **text**
        parts = re.split(r"(\*\*.*?\*\*)", text)
        p = doc.add_paragraph()
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## "):
            flush_para()
            heading_text = stripped[3:].strip()
            doc.add_heading(heading_text, level=2)
        elif stripped == "":
            flush_para()
        else:
            current_para_lines.append(stripped)

    flush_para()

    if page_break:
        doc.add_page_break()


def add_simple_section(doc: Document, title: str, text: str) -> None:
    """Add a short section with heading + single paragraph."""
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)


def add_table1(doc: Document, paper_dir: Path) -> None:
    """Parse table1.md and add as a Word table."""
    doc.add_heading("Table 1: Validation Gate Summary", level=2)

    raw = read_md(paper_dir / "table1.md")
    # Find markdown table rows (lines starting with |)
    table_lines = [
        ln.strip()
        for ln in raw.split("\n")
        if ln.strip().startswith("|") and not ln.strip().startswith("|--")
    ]
    if not table_lines:
        doc.add_paragraph("[Table 1 could not be parsed]")
        return

    # Parse rows
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.split("|")]
        cells = [c for c in cells if c != ""]  # remove empty edge cells
        rows.append(cells)

    if len(rows) < 2:
        doc.add_paragraph("[Table 1 has insufficient rows]")
        return

    # Create Word table
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols, style="Table Grid")

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    # Notes below table
    notes_match = re.search(r"\*Notes:\*(.+)", raw, re.DOTALL)
    if notes_match:
        notes_text = notes_match.group(1).strip()
        notes_text = re.sub(r"\s*\n\s*", " ", notes_text)
        p = doc.add_paragraph()
        run = p.add_run(notes_text)
        run.font.size = Pt(9)
        run.font.italic = True


def main() -> None:
    paper_dir = Path(r"D:\ipf_sprint\results\paper")
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Title Page ──
    add_title_page(doc)

    # ── Abstract ──
    add_abstract(doc, paper_dir)

    # ── Methods ──
    add_markdown_section(doc, "Methods", paper_dir / "methods_draft.md")

    # ── Results ──
    add_markdown_section(doc, "Results", paper_dir / "results_draft.md", page_break=False)

    # Table 1 (inline after Results)
    add_table1(doc, paper_dir)
    doc.add_page_break()

    # ── Discussion ──
    add_markdown_section(doc, "Discussion", paper_dir / "discussion_draft.md")

    # ── Data Availability ──
    data_text = read_md(paper_dir / "data_availability.md")
    data_text = re.sub(r"^##?\s+[^\n]+\n", "", data_text, count=1)
    data_text = re.sub(r"\*\*(.+?)\*\*", r"\1", data_text)
    data_text = data_text.strip()
    doc.add_heading("Data Availability", level=1)
    for para in data_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # ── Competing Interests ──
    add_simple_section(doc, "Competing Interests",
                       "The author declares no competing interests.")

    # ── Author Contributions ──
    add_simple_section(doc, "Author Contributions",
                       "TK conceived and designed the study, developed the validation "
                       "framework, performed all analyses, and wrote the manuscript.")

    # ── Acknowledgments ──
    add_simple_section(doc, "Acknowledgments",
                       "Pipeline development was assisted by AI tools (Claude 3.7 Sonnet, "
                       "GitHub Copilot) for code optimization and documentation.")

    # ── Figure Legends ──
    doc.add_page_break()
    doc.add_heading("Figure Legends", level=1)

    doc.add_heading("Figure 1", level=2)
    doc.add_paragraph(
        "Phase 2 validation: Baseline pathway Separability Index (SI) compared to "
        "matched random gene set controls. Histogram shows distribution of random SI "
        "(gray), with baseline SI indicated by vertical line (blue). Both cohorts show "
        "baseline significantly exceeding random controls (p < 0.001). "
        "Left panel: GSE24206 (baseline SI = 0.381, random mean = 0.225). "
        "Right panel: GSE53845 (baseline SI = 0.349, random mean = 0.221)."
    )

    doc.add_heading("Figure 2", level=2)
    doc.add_paragraph(
        "Phase 3 holdout validation: Subject-level generalization testing. Bar plots "
        "show holdout SI (colored) compared to permutation null (gray) with error bars "
        "indicating standard deviation across 50 random splits. "
        "Left panel: GSE24206 demonstrates significant generalization "
        "(SI = 0.499 +/- 0.136, p = 0.03, green). "
        "Right panel: GSE53845 shows borderline significance "
        "(SI = 0.467 +/- 0.139, p = 0.07, gold) with robust effect size "
        "(Delta = 0.164 vs. null mean 0.303)."
    )

    # ── Save ──
    output = paper_dir / "manuscript_combined.docx"
    doc.save(str(output))
    size_kb = output.stat().st_size / 1024
    print(f"Manuscript saved: {output}")
    print(f"File size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()

